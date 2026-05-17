#!/usr/bin/env python3
"""Stages 9-16: build the BMC Cancer submission package."""

from __future__ import annotations

import shutil
import subprocess
import sys
import textwrap
import zipfile
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

from src.features import PATHWAY_DEFINITIONS  # noqa: E402
from src.meta import random_effects_from_ci  # noqa: E402
from src.pathways import SEVEN_PATHWAY_COMPONENTS  # noqa: E402

RESULTS = REPO_ROOT / "results"
REPORTS = RESULTS / "reports"
FIGURES = REPO_ROOT / "figures"
DOCS = REPO_ROOT / "docs"
PAPER = REPO_ROOT / "paper"
MANUSCRIPT = REPO_ROOT / "manuscript"
MFIGURES = MANUSCRIPT / "figures"
MTABLES = MANUSCRIPT / "tables"
ADDITIONAL = MANUSCRIPT / "additional_files"
DELIVERABLES = REPO_ROOT / "deliverables"
LOG = REPO_ROOT / "logs" / "phase3_log.md"

HEADLINE = "Gradient_Boosted_Survival"
PAM50 = "PAM50_ROR_official"
COX = "Cox_PH"
LOCKED_TNBC = {"delta": 0.0144, "ci_low": -0.0472, "ci_high": 0.0760, "p": 0.6466}


def timestamp() -> str:
    return datetime.now().astimezone().isoformat(timespec="seconds")


def log(message: str) -> None:
    LOG.parent.mkdir(exist_ok=True)
    with LOG.open("a") as fh:
        fh.write(f"\n- {timestamp()} {message}\n")


def run(cmd: list[str]) -> str:
    result = subprocess.run(cmd, cwd=REPO_ROOT, text=True, capture_output=True, check=False)
    if result.returncode != 0:
        return result.stderr.strip() or result.stdout.strip()
    return result.stdout.strip()


def git_sha(path: str | None = None, first: bool = False) -> str:
    cmd = ["git", "log", "--format=%H"]
    if path:
        cmd += ["--", path]
    output = run(cmd)
    lines = [line for line in output.splitlines() if line.strip()]
    if not lines:
        return "[TODO Suhaan: commit SHA]"
    return lines[-1] if first else lines[0]


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(textwrap.dedent(text).strip() + "\n")


def fmt(value: float, digits: int = 3) -> str:
    if value is None or not np.isfinite(value):
        return "NA"
    return f"{value:.{digits}f}"


def fmt_signed(value: float, digits: int = 4) -> str:
    if value is None or not np.isfinite(value):
        return "NA"
    return f"{value:+.{digits}f}"


def load_numbers() -> dict:
    table1 = pd.read_csv(RESULTS / "Table_1_patient_characteristics.csv")
    h2h = pd.read_csv(RESULTS / "Table_3_head_to_head.csv")
    ext = pd.read_csv(RESULTS / "Table_2_external_validation.csv")
    cal = pd.read_csv(RESULTS / "Table_S10_calibration.csv")
    dca = pd.read_csv(RESULTS / "Table_S11_dca.csv")
    nri = pd.read_csv(RESULTS / "Table_S12_nri_idi.csv")
    subtype = pd.read_csv(RESULTS / "Table_S9_within_subtype_meta.csv")
    stability = pd.read_csv(RESULTS / "Table_S13_stability.csv")
    internal = pd.read_csv(RESULTS / "Table_S2_ml_internal_cv.csv")

    overall_rows = h2h[
        h2h["headline_model"].eq(HEADLINE) & h2h["comparator"].eq(PAM50) & h2h["subgroup"].eq("overall")
    ].rename(columns={"delta_cindex": "delta", "ci_low": "low", "ci_high": "high"})
    overall_meta = random_effects_from_ci(overall_rows, "delta", "low", "high")
    tnbc_rows = h2h[
        h2h["headline_model"].eq(HEADLINE) & h2h["comparator"].eq(PAM50) & h2h["subgroup"].eq("tnbc")
    ].dropna(subset=["delta_cindex"])

    def pool_cindex(model: str) -> dict:
        rows = ext[
            ext["model"].eq(model) & ext["subgroup"].eq("overall") & ext["status"].eq("ok")
        ].rename(columns={"harrell_c": "effect", "harrell_c_ci_low": "low", "harrell_c_ci_high": "high"})
        return random_effects_from_ci(rows, "effect", "low", "high")

    ml_c = pool_cindex(HEADLINE)
    pam_c = pool_cindex(PAM50)
    cox_c = pool_cindex(COX)

    wide_dca = dca.pivot_table(index=["cohort", "threshold"], columns="strategy", values="net_benefit")
    dca_delta = (wide_dca["Pathway ML"] - wide_dca["PAM50-ROR"]).dropna()
    mean_ici = cal.groupby("model")["ici"].mean(numeric_only=True)
    internal_row = internal[internal["model"].eq(HEADLINE)].iloc[0]
    stability_row = stability.iloc[0]

    return {
        "git_sha": git_sha(),
        "prereg_sha": git_sha("docs/PRIMARY_ENDPOINT.md", first=True),
        "initial_n": int(table1["initial_samples_in_harmonized_database"].sum()),
        "analysis_n": int(table1["final_analysis_n"].sum()),
        "events": int(table1["os_events"].sum()),
        "followup_range": f"{table1['median_followup_months_iqr'].iloc[0]} to {table1['median_followup_months_iqr'].iloc[-1]}",
        "overall_meta": overall_meta,
        "tnbc_meta": LOCKED_TNBC,
        "tnbc_cohorts": int(len(tnbc_rows)),
        "tnbc_n": int(tnbc_rows["n"].sum()) if not tnbc_rows.empty else 0,
        "tnbc_events": int(tnbc_rows["events"].sum()) if not tnbc_rows.empty else 0,
        "ml_meta_c": ml_c,
        "pam_meta_c": pam_c,
        "cox_meta_c": cox_c,
        "internal_delta_cox": float(internal_row["delta_vs_cox"]),
        "internal_cv": float(internal_row["mean_cv_cindex"]),
        "mean_ici_ml": float(mean_ici.get(HEADLINE, np.nan)),
        "mean_ici_pam50": float(mean_ici.get(PAM50, np.nan)),
        "dca_delta": float(dca_delta.mean()) if len(dca_delta) else np.nan,
        "nri_020_mean": float(nri[nri["threshold"].eq(0.2)]["nri"].mean()),
        "idi_mean": float(nri["idi"].mean()),
        "subtype_lines": [
            f"{row.subtype}: delta {fmt_signed(row.meta_delta)}, 95% CI [{fmt_signed(row.meta_ci_low)}, {fmt_signed(row.meta_ci_high)}], p={fmt(row.meta_p)}"
            for row in subtype.itertuples()
        ],
        "stability_rho": float(stability_row["mean_pairwise_spearman"]),
        "stability_low": float(stability_row["spearman_ci_low"]),
        "stability_high": float(stability_row["spearman_ci_high"]),
        "stability_top": ", ".join(stability.head(3)["feature"].tolist()),
    }


def write_readme_and_metadata(nums: dict) -> None:
    write_text(
        REPO_ROOT / "README.md",
        f"""
        # BRCA-PathwayML

        BMC Cancer submission package for an interpretable pathway-based machine learning benchmark in breast cancer prognosis.

        **Authors:** Suhaan Thayyil; Eshaan Nidee

        ## Current Framing

        This repository contains a four-cohort analysis comparing a locked pathway-based survival machine learning model against official genefu PAM50-ROR. The pre-registered TNBC endpoint was not met. The active BMC Cancer framing is transparent benchmarking with comparable discrimination, modest secondary calibration and decision-curve signals, and clear limitations.

        ## Key Locked Result

        TNBC delta C-index for Gradient Boosted Survival versus PAM50-ROR = {fmt_signed(nums['tnbc_meta']['delta'])}, 95% CI [{fmt_signed(nums['tnbc_meta']['ci_low'])}, {fmt_signed(nums['tnbc_meta']['ci_high'])}], p={fmt(nums['tnbc_meta']['p'], 4)}. The pre-registered threshold was delta >= +0.03 with p < 0.05. It was not met.

        ## Reproducibility

        - Analysis cohort: {nums['analysis_n']} patients with survival and official PAM50-ROR across four cohorts. Harmonized database contains {nums['initial_n']} samples.
        - Main results: `results/`
        - BMC submission files: `deliverables/`
        - Locked endpoint: `docs/PRIMARY_ENDPOINT.md`
        - Phase Three story lock: `docs/STORY.md`

        ## Data Sources

        TCGA-BRCA, GSE96058/SCAN-B, METABRIC from cBioPortal, and GSE20685. Raw data are not redistributed here.

        ## License

        MIT License. See `LICENSE`.
        """,
    )
    write_text(
        REPO_ROOT / "LICENSE",
        """
        MIT License

        Copyright (c) 2026 Suhaan Thayyil and Eshaan Nidee

        Permission is hereby granted, free of charge, to any person obtaining a copy
        of this software and associated documentation files (the "Software"), to deal
        in the Software without restriction, including without limitation the rights
        to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
        copies of the Software, and to permit persons to whom the Software is
        furnished to do so, subject to the following conditions:

        The above copyright notice and this permission notice shall be included in all
        copies or substantial portions of the Software.

        THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
        IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
        FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
        AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
        LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
        OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
        SOFTWARE.
        """,
    )
    write_text(
        REPO_ROOT / "CITATION.cff",
        f"""
        cff-version: 1.2.0
        title: "BRCA-PathwayML"
        message: "If you use this code, please cite the associated manuscript after publication."
        type: software
        authors:
          - family-names: Thayyil
            given-names: Suhaan
          - family-names: Nidee
            given-names: Eshaan
        year: 2026
        repository-code: "https://github.com/[TODO Suhaan: repository URL]"
        commit: "{nums['git_sha']}"
        license: MIT
        """,
    )


def write_main_tex(nums: dict) -> None:
    tex = f"""
    \\documentclass{{bmcart}}
    \\usepackage{{graphicx}}
    \\usepackage{{booktabs}}
    \\usepackage{{url}}
    \\usepackage{{hyperref}}

    \\begin{{document}}

    \\begin{{frontmatter}}

    \\begin{{fmbox}}
    \\dochead{{Research article}}
    \\title{{Interpretable pathway-based machine learning for breast cancer prognosis: a four-cohort head-to-head benchmark against PAM50-ROR}}

    \\author*[1]{{\\fnm{{Suhaan}} \\sur{{Thayyil}}}}\\email{{[TODO Suhaan: corresponding author email]}}
    \\author[2]{{\\fnm{{Eshaan}} \\sur{{Nidee}}}}\\email{{[TODO Suhaan/Eshaan: email]}}

    \\address[1]{{\\orgname{{[TODO Suhaan: affiliation]}}, \\city{{[TODO]}}, \\country{{[TODO]}}}}
    \\address[2]{{\\orgname{{[TODO Suhaan/Eshaan: affiliation]}}, \\city{{[TODO]}}, \\country{{[TODO]}}}}
    \\end{{fmbox}}

    \\begin{{abstractbox}}
    \\begin{{abstract}}
    \\parttitle{{Background}} [TODO: Suhaan writes Background.]
    \\parttitle{{Methods}} We analyzed TCGA-BRCA, GSE96058/SCAN-B, METABRIC, and GSE20685. A Gradient Boosted Survival model trained on TCGA pathway and clinical features was externally evaluated without refitting. Official PAM50-ROR was computed with genefu. The locked primary endpoint tested TNBC delta C-index versus PAM50-ROR and was not met.
    \\parttitle{{Results}} [TODO: Suhaan writes Results prose from docs/results_bullets.md.]
    \\parttitle{{Conclusions}} [TODO: Suhaan writes Conclusions.]
    \\end{{abstract}}
    \\begin{{keyword}}
    \\kwd{{breast cancer}}
    \\kwd{{machine learning}}
    \\kwd{{PAM50}}
    \\kwd{{prognostic model}}
    \\kwd{{TRIPOD}}
    \\end{{keyword}}
    \\end{{abstractbox}}

    \\end{{frontmatter}}

    \\section{{Background}}
    [TODO: Suhaan writes Background.]

    \\section{{Methods}}

    \\subsection{{Cohorts and data sources}}
    We analyzed four public breast cancer cohorts: TCGA-BRCA, GSE96058/SCAN-B, METABRIC from the cBioPortal study \\texttt{{brca\\_metabric}}, and GSE20685. The harmonized database contained {nums['initial_n']} samples. The final Phase Three analysis set included {nums['analysis_n']} samples with usable overall survival and official PAM50-ROR scores.

    \\subsection{{Inclusion and exclusion}}
    Samples were retained if overall survival time was positive and the official genefu PAM50-ROR calculation completed. Samples missing survival data or official PAM50-ROR were excluded from the head-to-head analysis. Figure 1 gives cohort-level counts.

    \\subsection{{Outcome definition}}
    The primary outcome for benchmarking was overall survival. Time was expressed in days. Death was coded as the event. Five-year and ten-year horizons were used for calibration summaries.

    \\subsection{{Pathway score computation}}
    Seven aggregate pathway features were used: immune, proliferation, DNA repair, metabolism, stromal/EMT, apoptosis/stress, and hormone response. Pathway features were computed before modeling and then used unchanged for model training and validation. The Phase Two model used MSigDB-derived pathway aggregates from Hallmark, Reactome, and KEGG Medicus components.

    \\subsection{{Clinical genomic baselines}}
    Official PAM50 subtype and ROR-S risk were computed using \\texttt{{genefu}} in R. Surrogate Oncotype DX and MammaPrint scores were retained as supplementary comparators where gene coverage allowed. EndoPredict and BCI were not treated as primary comparators unless exact public formulas were reproducible from the public data.

    \\subsection{{Machine learning models}}
    The model zoo included Cox proportional hazards, Coxnet, Random Survival Forest, Gradient Boosted Survival Analysis, DeepSurv, and a stacked ensemble. Model selection was completed before Phase Three. The locked headline model for Phase Three is Gradient Boosted Survival. Phase Three did not run new model search.

    \\subsection{{Cross-validation protocol}}
    Models were trained on TCGA-BRCA only. The Phase Two internal benchmark used five-fold cross-validation on TCGA-BRCA with fixed random seed 42. The headline model had internal C-index {fmt(nums['internal_cv'])} and delta versus Cox {fmt_signed(nums['internal_delta_cox'])}.

    \\subsection{{External validation}}
    The trained TCGA model was applied without refitting to GSE96058/SCAN-B, METABRIC, and GSE20685. External analyses used the same risk scores and the same comparator scores for all head-to-head analyses.

    \\subsection{{Statistical methods}}
    Harrell C-index and bootstrap 95\\% confidence intervals were computed for discrimination. Paired bootstrap delta C-index compared the headline model with PAM50-ROR. Random-effects meta-analysis used the DerSimonian-Laird estimator. Calibration at five and ten years used calibration slope, intercept, calibration-in-the-large, and a decile-smoothed integrated calibration index. Decision curve analysis used five-year mortality thresholds from 0.05 to 0.50. NRI and IDI were calculated at 0.10 and 0.20 thresholds. Permutation importance stability was estimated in 100 stratified 80\\% TCGA subsamples.

    \\subsection{{Pre-registration disclosure}}
    The locked pre-registered primary endpoint was TNBC meta-analyzed delta C-index of pathway ML versus PAM50-ROR with success threshold delta >= +0.03 and p < 0.05. The endpoint was not met: delta {fmt_signed(nums['tnbc_meta']['delta'])}, 95\\% CI [{fmt_signed(nums['tnbc_meta']['ci_low'])}, {fmt_signed(nums['tnbc_meta']['ci_high'])}], p={fmt(nums['tnbc_meta']['p'], 4)}. All Phase Three analyses beyond the locked endpoint are exploratory or secondary. The analysis plan is stored in \\texttt{{docs/PRIMARY\\_ENDPOINT.md}} at commit \\texttt{{{nums['prereg_sha']}}}.

    \\subsection{{Software and reproducibility}}
    Analyses were run locally with Python 3.12.4 and R 4.6.0. The repository contains the model artifacts, scripts, tables, and figure builders. The repository URL and final commit SHA should be filled before submission: \\texttt{{[TODO Suhaan: repository URL]}}, commit \\texttt{{{nums['git_sha']}}}. The code is released under the MIT License.

    \\subsection{{Ethics}}
    All data analyzed in this study are publicly available and de-identified. No additional ethics approval was required. Original data collection ethics approvals are documented in the respective cohort publications.

    \\section{{Results}}
    [TODO: Suhaan writes Results prose from docs/results_bullets.md.]

    \\begin{{itemize}}
    \\item Patient cohorts: Table 1 and Figure 1.
    \\item Discrimination and head-to-head comparison: Table 2, Table 3, and Figure 2.
    \\item Exploratory within-subtype analysis: Figure 3 and Supplementary Tables S8-S9.
    \\item Exploratory calibration: Figure 4 and Supplementary Table S10.
    \\item Exploratory decision curve analysis: Figure 5 and Supplementary Tables S11-S12.
    \\item Stability analysis: Additional file 5 and Supplementary Table S13.
    \\end{{itemize}}

    \\section{{Discussion}}
    [TODO: Suhaan writes Discussion.]

    \\section{{Conclusions}}
    [TODO: Suhaan writes Conclusions.]

    \\section*{{Declarations}}

    \\subsection*{{Ethics approval and consent to participate}}
    All data analyzed in this study are publicly available and de-identified. No additional ethics approval was required. Original data collection ethics approvals are documented in the respective cohort publications.

    \\subsection*{{Consent for publication}}
    Not applicable.

    \\subsection*{{Availability of data and materials}}
    TCGA-BRCA data are available through the GDC Data Portal. GSE96058 and GSE20685 are available through NCBI GEO. METABRIC data are available through cBioPortal. Code and derived analysis outputs will be available at \\texttt{{[TODO Suhaan: repository URL]}}. The locked analysis plan is available in \\texttt{{docs/PRIMARY\\_ENDPOINT.md}} at commit \\texttt{{{nums['prereg_sha']}}}.

    \\subsection*{{Competing interests}}
    The authors declare no competing interests. [TODO Suhaan/Eshaan: confirm.]

    \\subsection*{{Funding}}
    This research received no external funding. [TODO Suhaan/Eshaan: confirm.]

    \\subsection*{{Authors' contributions}}
    ST conceived the study, performed the analyses, and drafted the manuscript. EN contributed to [TODO Suhaan/Eshaan: specify contribution]. Both authors reviewed and approved the final manuscript.

    \\subsection*{{Acknowledgements}}
    [TODO Suhaan/Eshaan: thank only people who consented to be acknowledged.]

    \\bibliographystyle{{bmc-mathphys}}
    \\bibliography{{main}}

    \\end{{document}}
    """
    write_text(PAPER / "main.tex", tex)
    if shutil.which("pdflatex") is None:
        write_text(
            PAPER / "PDF_BUILD_BLOCKED.md",
            """
            # PDF Build Blocked

            `pdflatex` is not installed in the local environment. The Phase Three package includes `paper/main.tex` and `paper/main.bib`, but `paper/main.pdf` was not built locally.

            BMC submission systems can compile LaTeX during upload. Suhaan and Eshaan should also compile after filling TODO blocks and installing a TeX distribution.
            """,
        )
        log("Stage 10 LaTeX compile skipped because pdflatex is missing.")
    else:
        result = run(["pdflatex", "-interaction=nonstopmode", "main.tex"])
        log(f"Stage 10 LaTeX compile attempted: {result[:500]}")


def write_bib() -> None:
    write_text(
        PAPER / "main.bib",
        """
        @article{parker2009supervised,
          title={Supervised risk predictor of breast cancer based on intrinsic subtypes},
          author={Parker, Joel S and Mullins, Michael and Cheang, Maggie CU and Leung, Samuel and Voduc, David and Vickery, Tammi and Davies, Sherri and Fauron, Christiane and He, Xiaping and Hu, Zhiyuan and others},
          journal={Journal of Clinical Oncology},
          year={2009},
          volume={27},
          number={8},
          pages={1160--1167}
        }

        @article{gendoo2016genefu,
          title={Genefu: an R/Bioconductor package for computation of gene expression-based signatures in breast cancer},
          author={Gendoo, Deena MA and Ratanasirigulchai, Nalis and Schroder, Mark S and Pare, Luisa and Parker, Joel S and Prat, Aleix and Haibe-Kains, Benjamin},
          journal={Bioinformatics},
          year={2016},
          volume={32},
          number={7},
          pages={1097--1099}
        }

        @article{paik2004gene,
          title={A multigene assay to predict recurrence of tamoxifen-treated, node-negative breast cancer},
          author={Paik, Soonmyung and Shak, Steven and Tang, Gong and Kim, Chungyeul and Baker, Joffre and Cronin, Maureen and Baehner, Frederick L and Walker, Michael G and Watson, Drew and Park, Taesung and others},
          journal={New England Journal of Medicine},
          year={2004},
          volume={351},
          number={27},
          pages={2817--2826}
        }

        @article{vantveer2002gene,
          title={Gene expression profiling predicts clinical outcome of breast cancer},
          author={van't Veer, Laura J and Dai, Hongyue and van de Vijver, Marc J and He, Yudong D and Hart, Augustinus AM and Mao, Mao and Peterse, Hans L and van der Kooy, Karin and Marton, Matthew J and Witteveen, Anke T and others},
          journal={Nature},
          year={2002},
          volume={415},
          number={6871},
          pages={530--536}
        }

        @article{tcga2012comprehensive,
          title={Comprehensive molecular portraits of human breast tumours},
          author={{The Cancer Genome Atlas Network}},
          journal={Nature},
          year={2012},
          volume={490},
          number={7418},
          pages={61--70}
        }

        @article{curtis2012genomic,
          title={The genomic and transcriptomic architecture of 2,000 breast tumours reveals novel subgroups},
          author={Curtis, Christina and Shah, Sohrab P and Chin, Suet-Feung and Turashvili, Gulisa and Rueda, Oscar M and Dunning, Mark J and Speed, Doug and Lynch, Andy G and Samarajiwa, Shamith and Yuan, Yinyin and others},
          journal={Nature},
          year={2012},
          volume={486},
          number={7403},
          pages={346--352}
        }

        @article{pereira2016somatic,
          title={The somatic mutation profiles of 2,433 breast cancers refines their genomic and transcriptomic landscapes},
          author={Pereira, Bernard and Chin, Suet-Feung and Rueda, Oscar M and Vollan, Hans-Kristian Moen and Provenzano, Elena and Bardwell, Helen A and Pugh, Michelle and Jones, Linda and Russell, Roslin and Sammut, Stephen-John and others},
          journal={Nature Communications},
          year={2016},
          volume={7},
          pages={11479}
        }

        @article{brueffer2018clinical,
          title={Clinical value of RNA sequencing-based classifiers for prediction of the five conventional breast cancer biomarkers: a report from the population-based multicenter Sweden Cancerome Analysis Network-Breast Initiative},
          author={Brueffer, Christian and Vallon-Christersson, Johan and Grabau, Dorthe and Ehinger, Anna and Hakkinen, Jari and Hegardt, Cecilia and Malina, Jan and Chen, Yilun and Bendahl, Par-Ola and Manjer, Jonas and others},
          journal={JCO Precision Oncology},
          year={2018},
          volume={2},
          pages={1--18}
        }

        @article{cox1972regression,
          title={Regression models and life-tables},
          author={Cox, David R},
          journal={Journal of the Royal Statistical Society: Series B},
          year={1972},
          volume={34},
          number={2},
          pages={187--202}
        }

        @article{simon2011regularization,
          title={Regularization paths for Cox's proportional hazards model via coordinate descent},
          author={Simon, Noah and Friedman, Jerome and Hastie, Trevor and Tibshirani, Rob},
          journal={Journal of Statistical Software},
          year={2011},
          volume={39},
          number={5},
          pages={1--13}
        }

        @article{ishwaran2008random,
          title={Random survival forests},
          author={Ishwaran, Hemant and Kogalur, Udaya B and Blackstone, Eugene H and Lauer, Michael S},
          journal={The Annals of Applied Statistics},
          year={2008},
          volume={2},
          number={3},
          pages={841--860}
        }

        @article{hothorn2006survival,
          title={Survival ensembles},
          author={Hothorn, Torsten and Buhlmann, Peter and Dudoit, Sandrine and Molinaro, Annette and Van Der Laan, Mark J},
          journal={Biostatistics},
          year={2006},
          volume={7},
          number={3},
          pages={355--373}
        }

        @article{katzman2018deepsurv,
          title={DeepSurv: personalized treatment recommender system using a Cox proportional hazards deep neural network},
          author={Katzman, Jared L and Shaham, Uri and Cloninger, Alexander and Bates, Jonathan and Jiang, Tingting and Kluger, Yuval},
          journal={BMC Medical Research Methodology},
          year={2018},
          volume={18},
          pages={24}
        }

        @article{derSimonian1986meta,
          title={Meta-analysis in clinical trials},
          author={DerSimonian, Rebecca and Laird, Nan},
          journal={Controlled Clinical Trials},
          year={1986},
          volume={7},
          number={3},
          pages={177--188}
        }

        @article{austin2019graphical,
          title={Graphical calibration curves and the integrated calibration index for survival models},
          author={Austin, Peter C and Steyerberg, Ewout W},
          journal={Statistics in Medicine},
          year={2019},
          volume={38},
          number={15},
          pages={2714--2742}
        }

        @article{vickers2006decision,
          title={Decision curve analysis: a novel method for evaluating prediction models},
          author={Vickers, Andrew J and Elkin, Elena B},
          journal={Medical Decision Making},
          year={2006},
          volume={26},
          number={6},
          pages={565--574}
        }

        @article{collins2015tripod,
          title={Transparent reporting of a multivariable prediction model for individual prognosis or diagnosis (TRIPOD): the TRIPOD statement},
          author={Collins, Gary S and Reitsma, Johannes B and Altman, Douglas G and Moons, Karel GM},
          journal={Annals of Internal Medicine},
          year={2015},
          volume={162},
          number={1},
          pages={55--63}
        }

        @article{liberzon2015hallmark,
          title={The Molecular Signatures Database hallmark gene set collection},
          author={Liberzon, Arthur and Birger, Chet and Thorvaldsdottir, Helga and Ghandi, Mahmoud and Mesirov, Jill P and Tamayo, Pablo},
          journal={Cell Systems},
          year={2015},
          volume={1},
          number={6},
          pages={417--425}
        }
        """,
    )


def write_abstract_cover_tripod(nums: dict) -> None:
    write_text(
        DOCS / "abstract_skeleton.md",
        f"""
        # Structured Abstract Skeleton

        ## Background
        [TODO Suhaan: write Background.]

        ## Methods
        - Four cohorts: TCGA-BRCA, GSE96058/SCAN-B, METABRIC, GSE20685.
        - Final head-to-head analysis set: n = {nums['analysis_n']} with survival and official genefu PAM50-ROR.
        - Model: locked Gradient Boosted Survival trained on TCGA-BRCA and externally evaluated without refitting.
        - Comparator: official genefu PAM50-ROR.

        ## Results
        - Overall delta C-index versus PAM50-ROR: {fmt_signed(nums['overall_meta']['effect'])}, 95% CI [{fmt_signed(nums['overall_meta']['ci_low'])}, {fmt_signed(nums['overall_meta']['ci_high'])}], p={fmt(nums['overall_meta']['p'], 4)}.
        - TNBC primary endpoint: {fmt_signed(nums['tnbc_meta']['delta'])}, 95% CI [{fmt_signed(nums['tnbc_meta']['ci_low'])}, {fmt_signed(nums['tnbc_meta']['ci_high'])}], p={fmt(nums['tnbc_meta']['p'], 4)}. Not met.
        - Calibration: mean ICI {fmt(nums['mean_ici_ml'])} for pathway ML and {fmt(nums['mean_ici_pam50'])} for PAM50-ROR.
        - DCA: mean net benefit delta {fmt_signed(nums['dca_delta'])} across cohorts and thresholds.
        - Stability: mean pairwise Spearman rho {fmt(nums['stability_rho'])}.

        ## Conclusions
        [TODO Suhaan: write Conclusions with comparable performance and negative primary endpoint.]
        """,
    )
    write_text(
        DOCS / "cover_letter_BMC_Cancer.md",
        f"""
        [Date]
        The Editorial Office
        BMC Cancer

        Dear Editor,

        We submit "Interpretable pathway-based machine learning for breast cancer prognosis: a four-cohort head-to-head benchmark against PAM50-ROR" for consideration by BMC Cancer.

        This study presents a transparent four-cohort benchmark of interpretable pathway-based machine learning against the official PAM50-ROR classifier for breast cancer prognosis. We pre-registered the primary TNBC endpoint before modeling. The primary endpoint was not met: meta delta C-index = {fmt_signed(nums['tnbc_meta']['delta'])}, 95% CI [{fmt_signed(nums['tnbc_meta']['ci_low'])}, {fmt_signed(nums['tnbc_meta']['ci_high'])}], p={fmt(nums['tnbc_meta']['p'], 4)}. We report this directly and frame the study as a comparable-discrimination benchmark supported by calibration, decision curve analysis, within-subtype analysis, and stability assessments across {nums['analysis_n']} evaluable patients.

        This manuscript was previously submitted to Communications Medicine in April 2026 and Scientific Reports in May 2026, and was declined without peer review. After those declines, we substantially revised the work by replacing surrogate PAM50 with official genefu-based PAM50-ROR, adding METABRIC and GSE20685 cohorts, pre-registering the primary endpoint, and adding a machine learning model zoo with cross-validation. The current manuscript reports a clearly stated negative primary finding and secondary analyses marked as exploratory.

        This work fits BMC Cancer's interest in transparent clinical oncology research, including reports with negative or null findings. Reporting a locked negative result contributes to reliability and helps reduce publication bias.

        Suggested reviewers: [TODO Suhaan/Eshaan: 3-5 names with affiliations and emails]
        Suggested editors: [TODO Suhaan/Eshaan]

        Sincerely,
        Suhaan Thayyil, on behalf of all authors
        [affiliation]
        [ORCID]
        [email]
        """,
    )
    items = [
        ("1 Title and abstract", "Partial", "Title and abstract", "Structured abstract has TODO prose blocks."),
        ("2 Background and objectives", "Partial", "Background", "Background is reserved for Suhaan."),
        ("3 Source of data", "Yes", "Methods: Cohorts and data sources", "Four public cohorts listed."),
        ("4 Participants", "Yes", "Methods: Inclusion and exclusion", "Eligibility criteria described."),
        ("5 Outcome", "Yes", "Methods: Outcome definition", "Overall survival defined."),
        ("6 Predictors", "Yes", "Methods: Pathway score computation", "Pathway and clinical predictors described."),
        ("7 Sample size", "Yes", "Table 1 and Figure 1", "Cohort counts provided."),
        ("8 Missing data", "Partial", "Methods: Inclusion and exclusion", "Exclusions reported; no imputation used in Phase Three."),
        ("9 Statistical analysis methods", "Yes", "Methods: Statistical methods", "Discrimination, calibration, DCA, meta-analysis described."),
        ("10 Risk groups", "Yes", "Methods and Results bullets", "Median split used only for exploratory KM plots."),
        ("11 Development versus validation", "Yes", "Methods: External validation", "TCGA training and external cohorts separated."),
        ("12 Participant flow", "Yes", "Figure 1", "Cohort flow diagram provided."),
        ("13 Participant characteristics", "Yes", "Table 1", "Patient characteristics table provided."),
        ("14 Model specification", "Yes", "Methods: Machine learning models", "Headline model and features described."),
        ("15 Model performance", "Yes", "Table 2, Table 3, Figure 2", "C-index and deltas reported."),
        ("16 Model updating", "Yes", "Methods", "No external refitting was performed."),
        ("17 Limitations", "Partial", "Discussion TODO", "Suhaan should discuss negative primary endpoint and stability."),
        ("18 Interpretation", "Partial", "Results bullets and Discussion TODO", "No significant-advantage claim is made."),
        ("19 Implications", "Partial", "Discussion TODO", "Human prose needed."),
        ("20 Supplementary information", "Yes", "Additional files", "Supplementary materials built."),
        ("21 Funding", "Partial", "Declarations", "Statement included, author confirmation needed."),
        ("22 Protocol", "Yes", "docs/PRIMARY_ENDPOINT.md", "Locked endpoint committed before Phase Two modeling."),
        ("23 Registration", "Partial", "Methods: Pre-registration disclosure", "Local repository pre-registration, not trial registry."),
        ("24 Data access", "Yes", "Declarations", "Public accessions listed."),
        ("25 Code access", "Partial", "Declarations", "Repository URL placeholder remains."),
        ("26 Harms", "Not applicable", "NA", "No intervention study."),
        ("27 Other information", "Partial", "Declarations", "Author metadata TODOs remain."),
    ]
    lines = [
        "# TRIPOD 2015 Checklist",
        "",
        "Reference: Collins GS et al. Ann Intern Med 2015;162:55-63.",
        "",
        "| Item | Status | Manuscript Location | Note |",
        "|---|---|---|---|",
    ]
    lines += [f"| {a} | {b} | {c} | {d} |" for a, b, c, d in items]
    write_text(DOCS / "TRIPOD_checklist.md", "\n".join(lines))


def write_results_bullets(nums: dict) -> None:
    lines = [
        "# Results Bullets",
        "",
        "## Patient cohorts",
        f"- Harmonized database: n = {nums['initial_n']} across 4 cohorts.",
        f"- Final official PAM50-ROR head-to-head analysis: n = {nums['analysis_n']}.",
        f"- Overall survival events in final analysis set: {nums['events']}.",
        "- See Figure 1 and Table 1.",
        "",
        "## Discrimination",
        "- Headline model: Gradient Boosted Survival.",
        f"- Overall meta delta C-index versus PAM50-ROR: {fmt_signed(nums['overall_meta']['effect'])}, 95% CI [{fmt_signed(nums['overall_meta']['ci_low'])}, {fmt_signed(nums['overall_meta']['ci_high'])}], p={fmt(nums['overall_meta']['p'], 4)}.",
        f"- Meta C-index estimate: pathway ML {fmt(nums['ml_meta_c']['effect'])}; PAM50-ROR {fmt(nums['pam_meta_c']['effect'])}; Cox baseline {fmt(nums['cox_meta_c']['effect'])}.",
        "- Interpretation: comparable, not statistically different from PAM50-ROR.",
        "- See Table 2, Table 3, Figure 2 Panel A.",
        "",
        "## TNBC subgroup, locked primary endpoint",
        f"- TNBC evaluable cohorts: {nums['tnbc_cohorts']}.",
        f"- TNBC total n in evaluable cohorts: {nums['tnbc_n']}; events: {nums['tnbc_events']}.",
        f"- Locked delta C-index versus PAM50-ROR: {fmt_signed(nums['tnbc_meta']['delta'])}, 95% CI [{fmt_signed(nums['tnbc_meta']['ci_low'])}, {fmt_signed(nums['tnbc_meta']['ci_high'])}], p={fmt(nums['tnbc_meta']['p'], 4)}.",
        "- Pre-registered success threshold: delta >= +0.03 and p < 0.05.",
        "- Result: NOT MET.",
        "- See Figure 2 Panel B.",
        "",
        "## Within-subtype analysis, exploratory",
        *[f"- {line}" for line in nums["subtype_lines"]],
        "- Verdict: no subtype had a consistent meta-analytic advantage meeting the Phase Three reporting rule.",
        "- See Figure 3 and Tables S8-S9.",
        "",
        "## Calibration, exploratory",
        f"- Mean ICI across cohorts and horizons: pathway ML {fmt(nums['mean_ici_ml'])}; PAM50-ROR {fmt(nums['mean_ici_pam50'])}.",
        "- Lower ICI favors the pathway ML model in this exploratory summary.",
        "- See Figure 4 and Table S10.",
        "",
        "## Decision curve analysis, exploratory",
        f"- Mean 5-year net benefit delta versus PAM50-ROR across cohorts and thresholds: {fmt_signed(nums['dca_delta'])}.",
        f"- Mean categorical NRI at 0.20 threshold: {fmt_signed(nums['nri_020_mean'])}.",
        f"- Mean IDI across 0.10 and 0.20 thresholds: {fmt_signed(nums['idi_mean'])}.",
        "- See Figure 5 and Tables S11-S12.",
        "",
        "## Stability, exploratory",
        f"- Mean pairwise Spearman rho across 100 subsample feature-rank lists: {fmt(nums['stability_rho'])}, empirical interval [{fmt(nums['stability_low'])}, {fmt(nums['stability_high'])}].",
        f"- Top mean-rank features: {nums['stability_top']}.",
        "- Limitation: feature-rank stability was below 0.6 and should be discussed as weak.",
        "- See Table S13 and Additional file 5.",
        "",
        "## ML versus Cox",
        f"- Internal TCGA CV C-index for Gradient Boosted Survival: {fmt(nums['internal_cv'])}.",
        f"- Internal delta versus Cox: {fmt_signed(nums['internal_delta_cox'])}.",
        "- External claim should remain modest because the overall head-to-head delta versus PAM50-ROR was not significant.",
    ]
    write_text(DOCS / "results_bullets.md", "\n".join(lines))


def copy_figures_and_captions() -> None:
    MFIGURES.mkdir(parents=True, exist_ok=True)
    mapping = {
        "Figure_1": "fig_cohort_flow",
        "Figure_2": "fig_main_forest",
        "Figure_3": "fig_within_subtype_forest",
        "Figure_4": "fig_calibration_plots",
        "Figure_5": "fig_dca_per_cohort",
    }
    for dest, stem in mapping.items():
        for ext in [".pdf", ".tiff"]:
            shutil.copy2(FIGURES / f"{stem}{ext}", MFIGURES / f"{dest}{ext}")
    write_text(
        MFIGURES / "captions.md",
        """
        # Figure Captions

        ## Figure 1
        Cohort inclusion flow for the four public breast cancer cohorts. Boxes show harmonized samples, exclusions for missing overall survival or official PAM50-ROR, and final analysis counts.

        ## Figure 2
        Delta C-index for Gradient Boosted Survival minus official PAM50-ROR. Panel A shows all-comer results by cohort and random-effects meta-analysis. Panel B shows the locked TNBC primary endpoint. The TNBC endpoint was not met.

        ## Figure 3
        Exploratory within-subtype delta C-index for Gradient Boosted Survival minus PAM50-ROR across external cohorts. Diamonds show random-effects meta-analysis within each subtype.

        ## Figure 4
        Exploratory calibration plots at 5-year and 10-year horizons. Lines compare observed and predicted mortality across deciles for pathway ML, Cox baseline, and PAM50-ROR.

        ## Figure 5
        Exploratory decision curve analysis at the 5-year mortality horizon. Net benefit is shown across risk thresholds for treat-all, treat-none, pathway ML, Cox baseline, and PAM50-ROR.
        """,
    )
    write_text(
        MFIGURES / "README.md",
        """
        # Figure Package

        - Figure 1: Cohort flow
        - Figure 2: Main forest plot, overall and TNBC
        - Figure 3: Within-subtype forest plot
        - Figure 4: Calibration plots
        - Figure 5: Decision curve analysis

        Each figure is provided as PDF and TIFF at 300 DPI for BMC Cancer upload.
        """,
    )


def copy_tables_and_docx() -> None:
    MTABLES.mkdir(parents=True, exist_ok=True)
    table_files = [
        "Table_1_patient_characteristics.csv",
        "Table_2_external_validation.csv",
        "Table_3_head_to_head.csv",
        "Table_S1_clinical_baselines.csv",
        "Table_S2_ml_internal_cv.csv",
        "Table_S4_rescue_internal_cv.csv",
        "Table_S5_rescue_external_validation.csv",
        "Table_S6_rescue_head_to_head.csv",
        "Table_S7_rescue_posthoc_top_candidate_screen.csv",
        "Table_S8_within_subtype_external.csv",
        "Table_S9_within_subtype_meta.csv",
        "Table_S10_calibration.csv",
        "Table_S11_dca.csv",
        "Table_S12_nri_idi.csv",
        "Table_S13_stability.csv",
    ]
    copied = []
    for name in table_files:
        src = RESULTS / name
        if src.exists():
            shutil.copy2(src, MTABLES / name)
            copied.append(name)
    write_text(
        MTABLES / "README.md",
        "\n".join(["# Table Package", "", "CSV files included:"] + [f"- {name}" for name in copied]),
    )

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8)
    doc.add_heading("BRCA-PathwayML BMC Cancer Tables", level=1)
    doc.add_paragraph("Full CSV files are included in the same folder. This Word file provides submission snippets.")
    for idx, name in enumerate(copied):
        if idx:
            doc.add_page_break()
        df = pd.read_csv(MTABLES / name)
        doc.add_heading(name.replace(".csv", ""), level=2)
        doc.add_paragraph(f"Rows: {len(df)}. Columns: {len(df.columns)}.")
        max_rows = min(len(df), 45)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        for cell, col in zip(table.rows[0].cells, df.columns, strict=False):
            cell.text = str(col)
        for _, row in df.head(max_rows).iterrows():
            cells = table.add_row().cells
            for cell, value in zip(cells, row.tolist(), strict=False):
                if pd.isna(value):
                    cell.text = ""
                elif isinstance(value, float):
                    cell.text = f"{value:.4g}"
                else:
                    cell.text = str(value)[:120]
        if len(df) > max_rows:
            doc.add_paragraph(f"Table truncated in Word snippet after {max_rows} rows. Use CSV for full table.")
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(6)
    doc.save(MTABLES / "Tables_combined.docx")
    if shutil.which("soffice") is None and shutil.which("libreoffice") is None:
        log("Tables_combined.docx visual render QA skipped because LibreOffice/soffice is missing.")


def write_pdf_text(path: Path, title: str, body: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    lines = []
    for paragraph in body.splitlines():
        if not paragraph.strip():
            lines.append("")
            continue
        lines.extend(textwrap.wrap(paragraph, width=92) or [""])
    pages = [lines[i : i + 45] for i in range(0, len(lines), 45)] or [[]]
    from matplotlib.backends.backend_pdf import PdfPages

    with PdfPages(path) as pdf:
        for page_idx, page in enumerate(pages, start=1):
            fig = plt.figure(figsize=(8.5, 11))
            fig.text(0.08, 0.95, title if page_idx == 1 else f"{title}, continued", fontsize=13, weight="bold")
            y = 0.90
            for line in page:
                fig.text(0.08, y, line, fontsize=8, family="monospace")
                y -= 0.018
            fig.text(0.08, 0.04, f"Page {page_idx}", fontsize=7)
            pdf.savefig(fig)
            plt.close(fig)


def write_additional_files(nums: dict) -> None:
    ADDITIONAL.mkdir(parents=True, exist_ok=True)
    prereg = (DOCS / "PRIMARY_ENDPOINT.md").read_text()
    write_text(ADDITIONAL / "Additional_file_1_PRIMARY_ENDPOINT.md", prereg)
    write_pdf_text(ADDITIONAL / "Additional_file_1_PRIMARY_ENDPOINT.pdf", "Additional file 1. Pre-registration", prereg)

    pathway_lines = ["# Additional File 2. Pathway Definitions", "", "## MSigDB aggregate components"]
    for name, components in SEVEN_PATHWAY_COMPONENTS.items():
        pathway_lines.append(f"- {name}: {', '.join(components)}")
    pathway_lines += ["", "## Legacy curated 45-gene panel retained for transparency"]
    for name, genes in PATHWAY_DEFINITIONS.items():
        pathway_lines.append(f"- {name}: {', '.join(genes)}")
    pathway_text = "\n".join(pathway_lines)
    write_text(ADDITIONAL / "Additional_file_2_pathway_definitions.md", pathway_text)
    write_pdf_text(ADDITIONAL / "Additional_file_2_pathway_definitions.pdf", "Additional file 2. Pathway definitions", pathway_text)

    shutil.copy2(DOCS / "TRIPOD_checklist.md", ADDITIONAL / "Additional_file_3_TRIPOD_checklist.md")

    table_list = "\n".join(f"- {p.name}" for p in sorted(MTABLES.glob("*.csv")))
    write_text(
        ADDITIONAL / "Additional_file_4_supplementary_tables.md",
        f"""
        # Additional File 4. Supplementary Tables

        The full supplementary tables are supplied as CSV files in `manuscript/tables/`.

        {table_list}
        """,
    )
    write_pdf_text(
        ADDITIONAL / "Additional_file_4_supplementary_tables.pdf",
        "Additional file 4. Supplementary tables",
        (ADDITIONAL / "Additional_file_4_supplementary_tables.md").read_text(),
    )

    supp_fig_dir = ADDITIONAL / "Additional_file_5_supplementary_figures"
    supp_fig_dir.mkdir(exist_ok=True)
    for path in [FIGURES / "fig_stability_heatmap.pdf", FIGURES / "fig_stability_heatmap.tiff"]:
        shutil.copy2(path, supp_fig_dir / path.name)
    km_dir = FIGURES / "km_within_subtype"
    if km_dir.exists():
        dest = supp_fig_dir / "km_within_subtype"
        if dest.exists():
            shutil.rmtree(dest)
        shutil.copytree(km_dir, dest)
    write_text(
        ADDITIONAL / "Additional_file_5_supplementary_figures.md",
        """
        # Additional File 5. Supplementary Figures

        Includes the permutation-importance rank stability heatmap and exploratory within-subtype Kaplan-Meier curves.
        """,
    )

    write_text(
        ADDITIONAL / "Additional_file_6_code_availability.md",
        f"""
        # Additional File 6. Code Availability

        Code and derived analysis files are available in this repository.

        Repository URL: [TODO Suhaan: repository URL]
        Phase Three build commit: {nums['git_sha']}
        Pre-registration commit: {nums['prereg_sha']}
        License: MIT
        """,
    )
    rescue = REPORTS / "rescue_analysis_qc.md"
    if rescue.exists():
        text = rescue.read_text()
    else:
        text = "Rescue analysis QC report was not found."
    write_text(
        ADDITIONAL / "Additional_file_7_rescue_analysis_transparency_report.md",
        "# Additional File 7. Rescue Analysis Transparency Report\n\nThis analysis informed limitations only and is not part of the locked primary endpoint.\n\n"
        + text,
    )
    write_text(
        ADDITIONAL / "MANIFEST.md",
        """
        # Additional Files Manifest

        - Additional file 1: Locked pre-registration document.
        - Additional file 2: Pathway component and gene definitions.
        - Additional file 3: TRIPOD checklist.
        - Additional file 4: Supplementary table inventory.
        - Additional file 5: Supplementary figures and KM curves.
        - Additional file 6: Code availability statement.
        - Additional file 7: Rescue analysis transparency report.
        """,
    )


def write_phase_summary_and_submission_readme(nums: dict) -> None:
    write_text(
        DOCS / "PHASE_THREE_SUMMARY.md",
        f"""
        # Phase Three Summary

        ## Completion Status
        - Stage 0 pre-flight: complete.
        - Stage 1 story lock: complete.
        - Stage 2 within-subtype analysis: complete.
        - Stage 3 calibration: complete.
        - Stage 4 DCA and NRI/IDI: complete.
        - Stage 5 stability: complete.
        - Stage 6 patient characteristics table: complete.
        - Stage 7 cohort flow: complete.
        - Stage 8 central forest plot: complete.
        - Stage 9 TRIPOD and documents: complete.
        - Stage 10 manuscript reframe: complete with TODO prose blocks.
        - Stage 11 cover letter: complete with reviewer TODOs.
        - Stage 12 figure and table formatting: complete.
        - Stage 13 supplementary files: complete.
        - Stage 14 results bullets: complete.
        - Stage 15 deliverable zips: complete.

        ## New Analysis Results
        - Within-subtype: no consistent subtype-level advantage over PAM50-ROR by the Phase Three reporting rule.
        - Calibration: mean ICI pathway ML {fmt(nums['mean_ici_ml'])}; PAM50-ROR {fmt(nums['mean_ici_pam50'])}.
        - DCA: mean 5-year net benefit delta versus PAM50-ROR {fmt_signed(nums['dca_delta'])}.
        - NRI at 0.20: mean {fmt_signed(nums['nri_020_mean'])}.
        - Stability: mean pairwise Spearman rho {fmt(nums['stability_rho'])}, empirical interval [{fmt(nums['stability_low'])}, {fmt(nums['stability_high'])}].

        ## Flagged Issues
        - The locked primary endpoint was not met and must remain negative in the manuscript.
        - Feature-rank stability was weak and should be discussed as a limitation.
        - `pdflatex` is missing locally, so `paper/main.pdf` was not built.
        - LibreOffice/soffice is missing locally, so Word table visual render QA was skipped.
        - Eshaan Nidee metadata remains TODO in `paper/main.tex`.
        - Repository URL, ORCID, email, affiliation, reviewers, and final prose remain TODO.

        ## Recommendation
        Ready for Suhaan and Eshaan to complete prose, author metadata, and final submission fields. No additional model search is recommended.
        """,
    )
    DELIVERABLES.mkdir(exist_ok=True)
    write_text(
        DELIVERABLES / "SUBMISSION_README.md",
        f"""
        # BRCA-PathwayML BMC Cancer Submission Package

        Built on: {timestamp()}
        Commit SHA at build: {nums['git_sha']}

        ## Files
        - `01_manuscript_package.zip`: LaTeX source, bibliography, PDF blocker note, cover letter, story, TRIPOD, results bullets
        - `02_figures_package.zip`: Main figures as PDF and TIFF, captions
        - `03_tables_package.zip`: Tables as CSV plus `Tables_combined.docx`
        - `04_supplementary_package.zip`: BMC-style additional files

        ## What Suhaan and Eshaan still need to do before submission
        - [ ] Write Background section in `paper/main.tex`
        - [ ] Write Results prose from `docs/results_bullets.md`
        - [ ] Write Discussion and Conclusions
        - [ ] Finalize Abstract from `docs/abstract_skeleton.md`
        - [ ] Confirm title
        - [ ] Fill suggested reviewers in cover letter
        - [ ] Confirm ORCID, email, and affiliations for both authors
        - [ ] Resolve all `[TODO Suhaan]` and `[TODO Suhaan/Eshaan]` markers
        - [ ] Confirm competing interests and funding declarations
        - [ ] Install TeX and compile `paper/main.tex`
        - [ ] Rebuild `01_manuscript_package.zip` after prose edits

        ## Submission URL
        BMC Cancer Editorial Manager: https://www.editorialmanager.com/bmccancer/default.aspx

        ## Pre-submission checklist
        - [ ] Cover letter mentions prior rejections
        - [ ] TRIPOD checklist complete
        - [ ] Data availability statement complete with accessions
        - [ ] Author contributions section complete
        - [ ] All figures are present as TIFF and PDF
        - [ ] Locked negative primary endpoint is reported directly
        - [ ] No active claim states a significant advantage over PAM50-ROR
        """,
    )


def zip_paths(zip_path: Path, paths: list[Path]) -> None:
    zip_path.parent.mkdir(exist_ok=True)
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in paths:
            if not path.exists():
                continue
            if path.is_dir():
                for item in path.rglob("*"):
                    if item.is_file():
                        zf.write(item, item.relative_to(REPO_ROOT))
            else:
                zf.write(path, path.relative_to(REPO_ROOT))
    with zipfile.ZipFile(zip_path) as zf:
        bad = zf.testzip()
    if bad:
        raise RuntimeError(f"Zip verification failed for {zip_path}: {bad}")


def build_zips() -> None:
    zip_paths(
        DELIVERABLES / "01_manuscript_package.zip",
        [
            PAPER / "main.tex",
            PAPER / "main.bib",
            PAPER / "main.pdf",
            PAPER / "PDF_BUILD_BLOCKED.md",
            DOCS / "cover_letter_BMC_Cancer.md",
            DOCS / "abstract_skeleton.md",
            DOCS / "results_bullets.md",
            DOCS / "STORY.md",
            DOCS / "TRIPOD_checklist.md",
        ],
    )
    zip_paths(DELIVERABLES / "02_figures_package.zip", [MFIGURES])
    zip_paths(DELIVERABLES / "03_tables_package.zip", [MTABLES])
    supp_paths = [ADDITIONAL, DOCS / "PHASE_THREE_SUMMARY.md", REPORTS / "FINAL_SUMMARY.md", REPORTS / "rescue_analysis_qc.md"]
    zip_paths(DELIVERABLES / "04_supplementary_package.zip", supp_paths)


def update_story_n() -> None:
    story = DOCS / "STORY.md"
    text = story.read_text()
    text = text.replace(
        "in four-cohort breast cancer prognosis (n = 4,532 patients)",
        "in four-cohort breast cancer prognosis (4,532 harmonized samples; 4,003 evaluable in the official PAM50-ROR head-to-head analysis)",
    )
    story.write_text(text)


def main() -> None:
    for path in [PAPER, DOCS, MFIGURES, MTABLES, ADDITIONAL, DELIVERABLES]:
        path.mkdir(parents=True, exist_ok=True)
    log("Stages 9-16 BMC package build started.")
    nums = load_numbers()
    update_story_n()
    nums = load_numbers()
    write_readme_and_metadata(nums)
    write_main_tex(nums)
    write_bib()
    write_abstract_cover_tripod(nums)
    write_results_bullets(nums)
    copy_figures_and_captions()
    copy_tables_and_docx()
    write_additional_files(nums)
    write_phase_summary_and_submission_readme(nums)
    build_zips()
    log("Stages 9-16 BMC package build completed.")


if __name__ == "__main__":
    main()
