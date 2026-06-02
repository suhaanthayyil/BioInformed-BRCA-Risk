"""Build the reviewer-facing DOCX manuscript from the LaTeX source.

Pandoc converts the LaTeX title-block superscripts into OfficeMath objects that
render as square placeholder glyphs in LibreOffice. This script runs Pandoc and
then rewrites the title block with ordinary Word superscript runs.
"""

from __future__ import annotations

import argparse
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]


def _run_pandoc(tex_path: Path, bib_path: Path, docx_path: Path) -> None:
    subprocess.run(
        [
            "pandoc",
            str(tex_path),
            f"--bibliography={bib_path}",
            "--citeproc",
            "-o",
            str(docx_path),
        ],
        check=True,
        cwd=ROOT,
    )


def _clear_paragraph(paragraph) -> None:
    para = paragraph._p
    for child in list(para):
        if child.tag != qn("w:pPr"):
            para.remove(child)


def _set_font(run, size: int, *, bold: bool = False, superscript: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    if superscript:
        run.font.superscript = True


def _rewrite_title_block(docx_path: Path) -> None:
    doc = Document(docx_path)

    title = (
        "Pathway-based machine learning is competitive with but does not exceed "
        "PAM50-ROR for breast cancer prognosis: a pre-registered multi-cohort "
        "benchmark in 4,532 patients"
    )

    p0 = doc.paragraphs[0]
    _clear_paragraph(p0)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p0.add_run(title)
    _set_font(run, 16, bold=True)

    p1 = doc.paragraphs[1]
    _clear_paragraph(p1)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, superscript in [
        ("Suhaan Thayyil", False),
        ("1,*", True),
        (", Eshaan Nidee", False),
        ("2", True),
    ]:
        run = p1.add_run(text)
        _set_font(run, 14, bold=True, superscript=superscript)

    p2 = doc.paragraphs[2]
    _clear_paragraph(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliations = [
        "1 Marvin Ridge High School, Waxhaw, NC, USA",
        "2 Liberty High School, Frisco, TX, USA",
        "* Corresponding author: 3557971678@student.ucps.k12.nc.us",
        "ORCID: S.T. 0009-0003-4359-4400; E.N. 0009-0001-9881-6643",
    ]
    for index, line in enumerate(affiliations):
        run = p2.add_run(line)
        _set_font(run, 11)
        if index < len(affiliations) - 1:
            run.add_break()

    doc.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--tex", default="paper/main.tex")
    parser.add_argument("--bib", default="paper/main.bib")
    parser.add_argument("--out", default="paper/main.docx")
    args = parser.parse_args()

    tex_path = ROOT / args.tex
    bib_path = ROOT / args.bib
    docx_path = ROOT / args.out

    _run_pandoc(tex_path, bib_path, docx_path)
    _rewrite_title_block(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main()
